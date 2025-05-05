#!/usr/bin/env python
"""
Legal document processing tool using the Artemis framework.
"""

import argparse
import json
import logging
import os
import sys
import time
from pathlib import Path
from typing import Dict, List, Optional, Union, Any, Tuple

# Add parent directory to path to allow importing from artemis
sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))

from artemis.utils.config import load_config

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)

def parse_arguments():
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(description="Legal Document Analysis Tool")
    parser.add_argument(
        "--input", 
        type=str,
        required=True,
        help="Path to input document (PDF, DOCX, or TXT)"
    )
    parser.add_argument(
        "--output", 
        type=str,
        default=None,
        help="Path to output file (JSON format, defaults to input filename with .json extension)"
    )
    parser.add_argument(
        "--config", 
        type=str,
        default=str(Path(__file__).parent / "config.yaml"),
        help="Path to configuration YAML file"
    )
    parser.add_argument(
        "--model_path", 
        type=str,
        help="Path to pre-trained model (overrides config)"
    )
    parser.add_argument(
        "--jurisdiction",
        type=str,
        help="Legal jurisdiction (overrides config)"
    )
    parser.add_argument(
        "--analysis_type",
        type=str,
        default="full",
        choices=["full", "extraction", "risk", "summary"],
        help="Type of analysis to perform"
    )
    parser.add_argument(
        "--verbose",
        action="store_true",
        help="Enable verbose logging"
    )
    return parser.parse_args()

def load_document(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """
    Load a document from file and extract text.
    
    Args:
        file_path: Path to the document file
    
    Returns:
        Tuple of extracted text and document metadata
    """
    file_path = Path(file_path)
    file_extension = file_path.suffix.lower()
    
    logger.info(f"Loading document: {file_path}")
    
    if file_extension == '.pdf':
        return load_pdf(file_path)
    elif file_extension in ['.docx', '.doc']:
        return load_docx(file_path)
    elif file_extension == '.txt':
        return load_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")

def load_pdf(file_path: Path) -> Tuple[str, Dict[str, Any]]:
    """
    Load and extract text from a PDF file.
    
    Args:
        file_path: Path to PDF file
    
    Returns:
        Tuple of extracted text and document metadata
    """
    try:
        import pypdf
    except ImportError:
        logger.error("PyPDF library not found. Install with: pip install pypdf")
        sys.exit(1)
    
    try:
        pdf = pypdf.PdfReader(file_path)
        
        # Extract metadata
        metadata = {
            "title": pdf.metadata.get('/Title', ''),
            "author": pdf.metadata.get('/Author', ''),
            "subject": pdf.metadata.get('/Subject', ''),
            "creator": pdf.metadata.get('/Creator', ''),
            "producer": pdf.metadata.get('/Producer', ''),
            "creation_date": str(pdf.metadata.get('/CreationDate', '')),
            "modification_date": str(pdf.metadata.get('/ModDate', '')),
            "page_count": len(pdf.pages)
        }
        
        # Extract text
        text = ""
        for page_num, page in enumerate(pdf.pages, 1):
            page_text = page.extract_text()
            if page_text:
                text += f"\n\n--- Page {page_num} ---\n\n{page_text}"
        
        return text, metadata
    
    except Exception as e:
        logger.error(f"Error loading PDF: {e}")
        raise

def load_docx(file_path: Path) -> Tuple[str, Dict[str, Any]]:
    """
    Load and extract text from a DOCX file.
    
    Args:
        file_path: Path to DOCX file
    
    Returns:
        Tuple of extracted text and document metadata
    """
    try:
        from docx import Document
        from docx.opc.coreprops import CoreProperties
    except ImportError:
        logger.error("python-docx library not found. Install with: pip install python-docx")
        sys.exit(1)
    
    try:
        doc = Document(file_path)
        
        # Extract metadata
        try:
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or '',
                "author": core_props.author or '',
                "subject": core_props.subject or '',
                "keywords": core_props.keywords or '',
                "last_modified_by": core_props.last_modified_by or '',
                "created": str(core_props.created) if core_props.created else '',
                "modified": str(core_props.modified) if core_props.modified else '',
                "paragraph_count": len(doc.paragraphs),
                "section_count": len(doc.sections)
            }
        except:
            # Fallback if metadata extraction fails
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "section_count": len(doc.sections)
            }
        
        # Extract text
        text = "\n\n".join([para.text for para in doc.paragraphs if para.text])
        
        return text, metadata
    
    except Exception as e:
        logger.error(f"Error loading DOCX: {e}")
        raise

def load_txt(file_path: Path) -> Tuple[str, Dict[str, Any]]:
    """
    Load and extract text from a plain text file.
    
    Args:
        file_path: Path to text file
    
    Returns:
        Tuple of extracted text and document metadata
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # Basic metadata
        stat = file_path.stat()
        metadata = {
            "file_size": stat.st_size,
            "created": time.ctime(stat.st_ctime),
            "modified": time.ctime(stat.st_mtime),
            "line_count": text.count('\n') + 1
        }
        
        return text, metadata
    
    except Exception as e:
        logger.error(f"Error loading text file: {e}")
        raise

def initialize_analyzer(config_path: str, model_path: Optional[str] = None, jurisdiction: Optional[str] = None):
    """
    Initialize the legal document analyzer.
    
    Args:
        config_path: Path to configuration file
        model_path: Optional path to pre-trained model
        jurisdiction: Optional jurisdiction override
    
    Returns:
        Initialized analyzer
    """
    from examples.legal_analysis.analyzer import LegalDocumentAnalyzer
    
    # Load configuration
    config = load_config(config_path)
    
    # Override model path if provided
    if model_path:
        config['model']['model_path'] = model_path
    
    # Override jurisdiction if provided
    if jurisdiction:
        config['legal_extensions']['jurisdiction'] = jurisdiction
    
    # Initialize analyzer
    logger.info(f"Initializing legal document analyzer with {config['model']['base_model']}")
    analyzer = LegalDocumentAnalyzer(config)
    
    return analyzer

def analyze_document(analyzer, document_text: str, metadata: Dict[str, Any], analysis_type: str) -> Dict[str, Any]:
    """
    Analyze a legal document.
    
    Args:
        analyzer: Legal document analyzer
        document_text: Text content of the document
        metadata: Document metadata
        analysis_type: Type of analysis to perform
    
    Returns:
        Analysis results
    """
    logger.info(f"Performing {analysis_type} analysis")
    
    # Perform analysis based on type
    if analysis_type == "extraction":
        results = analyzer.extract_information(document_text, metadata)
    elif analysis_type == "risk":
        results = analyzer.analyze_risks(document_text, metadata)
    elif analysis_type == "summary":
        results = analyzer.generate_summary(document_text, metadata)
    else:  # "full"
        results = analyzer.analyze(document_text, metadata)
    
    return results

def save_results(results: Dict[str, Any], output_path: str):
    """
    Save analysis results to a file.
    
    Args:
        results: Analysis results
        output_path: Path to output file
    """
    # Create directory if it doesn't exist
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Save results
    logger.info(f"Saving results to {output_path}")
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=2, ensure_ascii=False)

def main():
    """Main function."""
    # Parse arguments
    args = parse_arguments()
    
    # Set logging level
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    try:
        # Load document
        document_text, metadata = load_document(args.input)
        
        # Set output path if not provided
        if args.output is None:
            input_path = Path(args.input)
            args.output = str(input_path.with_suffix('.json'))
        
        # Initialize analyzer
        analyzer = initialize_analyzer(args.config, args.model_path, args.jurisdiction)
        
        # Analyze document
        start_time = time.time()
        results = analyze_document(analyzer, document_text, metadata, args.analysis_type)
        end_time = time.time()
        
        # Add processing information
        results['processing_info'] = {
            'input_file': args.input,
            'analysis_type': args.analysis_type,
            'processing_time': f"{end_time - start_time:.2f} seconds"
        }
        
        # Save results
        save_results(results, args.output)
        
        logger.info(f"Analysis completed in {end_time - start_time:.2f} seconds")
        logger.info(f"Results saved to {args.output}")
    
    except Exception as e:
        logger.error(f"Error processing document: {e}", exc_info=True)
        sys.exit(1)

if __name__ == "__main__":
    main()
